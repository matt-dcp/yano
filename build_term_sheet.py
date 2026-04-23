#!/usr/bin/env python3
"""Build Casa Yano loan structure proposal (term sheet) as Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# Colors
CHARCOAL = RGBColor(0x2D, 0x2D, 0x2D)
GOLD = RGBColor(0xA6, 0x8A, 0x3E)
WARM_GRAY = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x2D, 0x8B, 0x2D)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)


def shade_cell(cell, color_hex):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_title(text, size=20, color=CHARCOAL, bold=True, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = 'Calibri'
    return p


def add_subtitle(text, size=11, color=WARM_GRAY, italic=False, space_after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.italic = italic
    return p


def add_section_heading(text, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text.upper())
    run.font.size = Pt(11)
    run.font.color.rgb = GOLD
    run.font.bold = True
    # Add bottom border
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C5A55A')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def add_body(text, size=10.5, bold=False, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p


def add_terms_table(rows):
    """Add a 2-column term sheet table (label | value)."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.0)
    tbl.columns[1].width = Inches(4.8)
    for i, (label, val) in enumerate(rows):
        tbl.rows[i].cells[0].width = Inches(2.0)
        tbl.rows[i].cells[1].width = Inches(4.8)
        c1 = tbl.rows[i].cells[0].paragraphs[0]
        c2 = tbl.rows[i].cells[1].paragraphs[0]
        c1.paragraph_format.space_after = Pt(2)
        c2.paragraph_format.space_after = Pt(2)
        r1 = c1.add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = c2.add_run(val)
        r2.font.size = Pt(10)
        if i % 2 == 0:
            shade_cell(tbl.rows[i].cells[0], "F5F5F5")
            shade_cell(tbl.rows[i].cells[1], "F5F5F5")
    return tbl


# ======================= HEADER =======================
add_title("Casa Yano", size=24, color=GOLD)
add_subtitle("210 W Yanonali Street, Santa Barbara, CA 93101", size=12, color=CHARCOAL)
add_subtitle("Loan Structure Proposal - Graduated Advance Facility", size=13, color=CHARCOAL, space_after=2)
add_subtitle("Prepared for Ben Scott, Montecito Bank & Trust   |   April 22, 2026", size=10, color=WARM_GRAY, italic=True, space_after=14)

# ======================= OVERVIEW =======================
add_section_heading("Structure Overview")
add_body("This proposal outlines a two-component loan structure: an initial funded advance at closing plus a "
         "committed earn-out facility that funds incrementally based on proven trailing operating performance. "
         "The structure captures the economic value of the property's strong early operating performance while "
         "staging credit exposure to proven results.", space_after=10)

# ======================= BASE LOAN =======================
add_section_heading("Base Loan (Funded at Closing)")
add_terms_table([
    ("Loan Amount", "$2,000,000"),
    ("Borrower", "DCP Wealth Fund, LLC (property-owning entity)"),
    ("Guarantors", "20%+ owners and control parties"),
    ("Rate", "5-Year Treasury + 2.35% spread (approx. 6.25% as of April 2026)"),
    ("Rate Term", "Held at approval; reset at Year 5 (spread remains fixed)"),
    ("Amortization", "25 years (30 years if residential classification applies)"),
    ("Loan Term", "10 years"),
    ("Loan Fee", "0.50% ($10,000)"),
    ("Prepayment Penalty", "5/5: 3-2-1-1-1 / 3-2-1-1-1"),
    ("Collateral", "First trust deed on real estate"),
    ("Coverage Requirement", "1.35x minimum DSCR at close"),
])

# ======================= EARN-OUT =======================
add_section_heading("Committed Earn-Out Facility")
add_terms_table([
    ("Maximum Facility Size", "$1,300,000 (committed at closing)"),
    ("Commitment Period", "24 months from closing"),
    ("Pricing", "Same spread (2.35%) over 5-Year Treasury at time of funding"),
    ("Upsize Fee", "0.25% on drawn amount only (no commitment fee on undrawn portion)"),
    ("Amortization", "Remaining term of base loan, recalibrated at each draw"),
    ("Coverage Requirement", "Minimum 1.40x DSCR on trailing period NOI at each draw"),
    ("Draw Type", "Borrower option (right, not obligation)"),
])

add_body("", space_after=4)
add_body("Tiered Funding Triggers:", bold=True, space_after=4)

# Triggers table with header
tbl = doc.add_table(rows=4, cols=3)
tbl.autofit = False
tbl.columns[0].width = Inches(1.8)
tbl.columns[1].width = Inches(2.8)
tbl.columns[2].width = Inches(2.2)

# Header row
header_cells = tbl.rows[0].cells
headers = ["Test Window", "Trigger", "Aggregate Earn-Out Available"]
for i, h in enumerate(headers):
    header_cells[i].width = tbl.columns[i].width
    p = header_cells[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(h)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(10)
    shade_cell(header_cells[i], "2D2D2D")

trigger_data = [
    ("T6 Annualized (Oct 2026)", "T6 annualized NOI >= $350,000", "Up to $500,000"),
    ("T12 Actual (Apr 2027)", "T12 actual NOI >= $375,000", "Up to $900,000 aggregate"),
    ("T12 Actual (Apr 2027)", "T12 actual NOI >= $425,000", "Up to $1,300,000 aggregate"),
]
for i, row_data in enumerate(trigger_data):
    row = tbl.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].width = tbl.columns[j].width
        p = row.cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(val)
        r.font.size = Pt(10)
        if j == 2:  # Aggregate column
            r.font.bold = True
        if i % 2 == 0:
            shade_cell(row.cells[j], "F5ECD7")

add_body("", space_after=4)
add_body("Each draw must independently satisfy the minimum 1.40x DSCR test. Earn-out rights expire if not drawn "
         "within the 24-month commitment period.", size=9, italic=True, space_after=10)

# ======================= BENEFITS TO LENDER =======================
add_section_heading("Benefits to MBT / Lender")

bullets_lender = [
    ("Staged credit exposure.", "Initial funding of $2.0M is based on in-place performance and stabilized appraisal. "
     "Incremental $1.3M funds only after proven trailing performance meets defined coverage thresholds. This is a "
     "structurally safer deployment of capital than writing the full amount on Day 1 against projected NOI."),
    ("Protects the banking relationship.", "Without an earn-out, the borrower is likely to seek supplemental debt from "
     "a non-relationship lender (DSCR or CMBS) to reach targeted proceeds, fragmenting the relationship and the deposit "
     "base. This structure keeps all debt - and the corresponding deposits, treasury, and follow-on opportunities - at MBT."),
    ("Credit-committee friendly.", "The structure presents as: 'We are lending $2M today on strong in-place performance, "
     "with optional upsize only after proven trailing NOI meets defined thresholds at defined coverage.' That framing "
     "is easier to approve than a speculative $3M loan based on projections."),
    ("Higher yield on proven credit.", "If performance delivers, MBT deploys additional capital at attractive spreads "
     "on a strengthened credit. If performance does not deliver, the capital is never at risk. Asymmetric upside for "
     "the bank."),
    ("Efficient marginal underwriting.", "The bank has already performed underwriting on this borrower and property. "
     "Incremental earn-out draws require marginal credit review against pre-defined triggers, not full re-underwriting. "
     "High origination ROI on the incremental proceeds."),
    ("Defensive against refinance risk.", "Without this structure, the borrower may refinance with another institution "
     "in 12-18 months to access proceeds MBT is not providing. The earn-out locks the borrower into a 10-year "
     "relationship at origination."),
]

for heading, text in bullets_lender:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(heading + " ")
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)

# ======================= COVERAGE ANALYSIS =======================
add_section_heading("Coverage Analysis at Each Funding Tier")

add_body("Based on the blended 2026 pro forma (closed actuals through March plus forward bookings and seasonal "
         "projections), projected full-year 2026 NOI after known fixed costs is $338,428. Assuming modest 5% "
         "annual growth into stabilization:", space_after=8)

# DSCR table
tbl2 = doc.add_table(rows=5, cols=4)
tbl2.autofit = False
tbl2.columns[0].width = Inches(1.7)
tbl2.columns[1].width = Inches(1.6)
tbl2.columns[2].width = Inches(1.6)
tbl2.columns[3].width = Inches(1.8)

hdr2 = ["Scenario", "Total Loan", "Annual DS (25-yr)", "DSCR @ $375K NOI"]
for i, h in enumerate(hdr2):
    tbl2.rows[0].cells[i].width = tbl2.columns[i].width
    p = tbl2.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(h)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(10)
    shade_cell(tbl2.rows[0].cells[i], "2D2D2D")

cov_data = [
    ("Base Loan Only", "$2,000,000", "$158,321", "2.37x"),
    ("Base + $500K Earn-Out", "$2,500,000", "$197,901", "1.89x"),
    ("Base + $900K Earn-Out", "$2,900,000", "$229,565", "1.63x"),
    ("Base + $1.3M Earn-Out", "$3,300,000", "$261,229", "1.44x"),
]
for i, row in enumerate(cov_data):
    r = tbl2.rows[i+1]
    for j, val in enumerate(row):
        r.cells[j].width = tbl2.columns[j].width
        p = r.cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(val)
        run.font.size = Pt(10)
        if j == 0 or j == 3:
            run.font.bold = True
        if i % 2 == 0:
            shade_cell(r.cells[j], "F5F5F5")

add_body("", space_after=4)
add_body("Even at the maximum earn-out, DSCR remains at 1.44x - comfortably above the 1.35x covenant with meaningful "
         "cushion for appraiser haircut, rate movement, or interim performance volatility.", size=9, italic=True, space_after=10)

# ======================= ALIGNMENT =======================
add_section_heading("Structural Alignment")

add_body("This structure aligns borrower and lender interests:", space_after=6)

alignment = [
    "Both parties benefit when the property continues to perform - MBT deploys more capital at attractive yields, "
    "borrower accesses capital for accretive redeployment.",
    "Both parties are protected if performance underperforms - MBT is not over-extended against projections, "
    "borrower is not forced to draw capital that cannot be covered.",
    "Both parties benefit from the long-term relationship captured by the structure - MBT retains a 10-year lending "
    "relationship and deposit base, borrower retains relationship-based pricing and flexibility.",
]
for a in alignment:
    add_bullet(a)

add_body("", space_after=10)

# ======================= NEXT STEPS =======================
add_section_heading("Proposed Next Steps")

next_steps = [
    "MBT review and feedback on proposed structure, triggers, and coverage thresholds",
    "Confirmation of pricing mechanics (spread lock, Treasury reset mechanics, earn-out pricing)",
    "Credit-committee presentation outline (borrower available to provide supporting materials as needed)",
    "Engagement of appraiser and environmental review following mutual agreement on structure",
]
for ns in next_steps:
    add_bullet(ns)

add_body("", space_after=20)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Proposal prepared by Driven Capital Partners   |   Indicative - subject to mutual agreement and credit approval")
r.font.size = Pt(8.5)
r.font.italic = True
r.font.color.rgb = WARM_GRAY

# Save
OUT = Path(__file__).parent / "Casa_Yano_Loan_Structure_Proposal.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
